# -*- coding: utf-8 -*-
"""
Created on Sun Mar 24 11:05:08 2024

@author: tevsl
"""

from typing import Optional, Any, Annotated, List, Dict, Literal


from pydantic import BaseModel, Field
import string
import re
import requests

def load_binary_file_from_url(url):
    try:
        # Send a GET request to the URL
        response = requests.get(url)

        # Check if the request was successful
        response.raise_for_status()

        # Access the binary content
        binary_content = response.content

        return binary_content

    except requests.exceptions.RequestException as e:
        print(f"An error occurred: {e}")
        return None
    
def find_direct_quotes(text):
    """
    Finds all direct quotes in a given text string, handling embedded apostrophes.

    Args:
        text (str): The input text string.

    Returns:
        list: A list of strings containing the direct quotes found in the text. 
              The list may be empty if no quotes are found.
    """
    # Regular expression to handle double quotes and single quotes with embedded apostrophes
    pattern = r'["“](.*?)["”]'
    
    # Find all matches
    matches = re.findall(pattern, text)
    
    return matches

def find_missing_strings(strings, transcript):
    """
    Checks whether each string in the input list is found in the transcript.
    Case and punctuation are ignored for the comparison.
    
    Args:
        strings (list): List of strings to check.
        transcript (str): The text of the transcript.
    
    Returns:
        list: Strings that are not found in the transcript.
    """
    # Normalize transcript by removing punctuation, converting to lower case, and removing spaces
    translator = str.maketrans('', '', string.punctuation)
    normalized_transcript = transcript.translate(translator).replace(" ", "").replace("\n","").lower()
    
    # Find strings not in the transcript
    missing_strings = []
    for s in strings:
        normalized_string = s.translate(translator).replace(" ", "").replace("\n","").lower()
        if normalized_string not in normalized_transcript:
            missing_strings.append(s)
    
    return missing_strings

def replace_special_characters(text):
    import re
    # Define the regex patterns and their replacements
    replacements = [
        (r'[“”]', '"'),      # Replace left and right double curly quotes with straight double quote
        (r'[‘’]', "'"),      # Replace left and right single curly quotes with straight single quote
        (r'—', '--'),        # Replace em dash with double hyphen
        (r'–', '-'),         # Replace en dash with hyphen
        (r'‐', '-')          # Replace non-standard hyphen with standard hyphen
    ]

    # Apply all replacements
    for pattern, replacement in replacements:
        text = re.sub(pattern, replacement, text)

    return text


def markdown_to_html(markdown_text, scrub_text=True,title=None, date=None, smart_transcript=None):
    """
    Convert markdown text to a self-contained HTML page with optional title, date, and smart transcript link.
    
    Args:
        markdown_text (str): The markdown content to be converted.
        title (str, optional): The title of the HTML document. Defaults to None.
        date (str, optional): The date for the document in 'YYYY-MM-DD' format. Defaults to None.
        smart_transcript (str, optional): URL to the smart transcript. Defaults to None.
        
    Returns:
        str: A self-contained HTML page as a string.
    """
    import markdown
    from datetime import datetime
    # Convert the markdown to HTML
    if scrub_text:
        markdown_text=replace_special_characters(markdown_text)
    html_content = markdown.markdown(markdown_text)
    
    # Format the title and date
    title_html = f"<h1>{title}</h1>\n" if title else ""
    date_html = f"<p class=\"date\">{date}</p>\n" if date else ""
    
    # Define CSS for styling
    css = """
    <style>
        body {
            font-family: 'Noto Sans', 'Arial', 'Helvetica', sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f4f4f4;
            color: #333;
        }
        h1 {
            color: #0056b3;
        }
        .date {
            font-size: 0.9em;
            color: #777;
        }
        p {
            margin: 1em 0;
        }
        a {
            color: #0056b3;
            text-decoration: none;
        }
        a:hover {
            text-decoration: underline;
        }
        pre {
            background: #333;
            color: #f8f8f2;
            padding: 10px;
            overflow-x: auto;
        }
        code {
            background: #f4f4f4;
            padding: 2px 4px;
        }
    </style>
    """
    
    # Format the smart transcript link
    smart_transcript_html = f"<p><a href=\"{smart_transcript}\" target=\"_blank\">See smart transcript</a></p>\n" if smart_transcript else ""
    
    # Construct the full HTML page
    html_template = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{title if title else 'Document'}</title>
        {css}
    </head>
    <body>
        {title_html}
        {date_html}
        {html_content}
        {smart_transcript_html}
    </body>
    </html>
    """ 
    return html_template

def generate_html_from_html(input_html_filepath, video_filepath, pk1_filepath, output_html_filepath=None):
   
    import pickle
    # Read the HTML from the input file
    with open(input_html_filepath, 'r', encoding='utf-8') as html_file:
        article_html = html_file.read()
    
    # Load the list of objects from the .pk1 file
    with open(pk1_filepath, 'rb') as pk1_file:
        deepgram_return = pickle.load(pk1_file)
        
    return generate_html_from_html_data(article_html, video_filepath, deepgram_return, output_html_filepath)

def generate_html_from_html_data(article_html, video_filepath, deepgram_return, output_html_filepath=None):
    import re
    from bs4 import BeautifulSoup
    from string import punctuation

    word_objects=deepgram_return['channels'][0]['alternatives'][0]['words']
    
    def strip_punctuation(word):
        return word.strip(punctuation).lower()
    
    def find_clip_for_quote(quote):
        # Replace hyphens with spaces and split into words
        words = [strip_punctuation(word) for word in quote.replace('-', ' ').split()]
        matches = []
        for i in range(len(word_objects) - len(words) + 1):
            if all(strip_punctuation(word_objects[i + j]['word']) == words[j] for j in range(len(words))):
                matches.append((word_objects[i]['start'], word_objects[i + len(words) - 1]['end']))
        return matches[0] if matches else None

    # Parse the HTML content
    soup = BeautifulSoup(article_html, 'html.parser')

    # Set the maximum column width and video player style
    style_tag = soup.new_tag('style')
    style_tag.string = """
    body {
        max-width: 7.5in;
        margin: auto;
    }
    video {
        display: block;
        width: 100%;
        height: auto;
        margin-bottom: 1em;
    }
    """
    soup.head.append(style_tag)

    # Determine if the file is an mp4 or m3u8
    is_m3u8 = video_filepath.lower().endswith('.m3u8')
    if is_m3u8:
        # Add hls.js script to the HTML
        script_tag = soup.new_tag('script', src='https://cdn.jsdelivr.net/npm/hls.js@latest')
        soup.head.append(script_tag)

        # Add custom script to initialize hls.js for each m3u8 video
        custom_script_tag = soup.new_tag('script')
        custom_script_tag.string = """
        document.addEventListener('DOMContentLoaded', function () {
            document.querySelectorAll('video[data-hls]').forEach(function (video) {
                var startTime = parseFloat(video.getAttribute('data-start'));
                var endTime = parseFloat(video.getAttribute('data-end'));

                if (Hls.isSupported()) {
                    var hls = new Hls();
                    hls.loadSource(video.getAttribute('data-hls'));
                    hls.attachMedia(video);
                    hls.on(Hls.Events.MANIFEST_PARSED, function () {
                        video.currentTime = startTime;
                        video.play();
                    });
                    video.addEventListener('timeupdate', function () {
                        if (video.currentTime >= endTime) {
                            video.pause();
                        }
                    });
                } else if (video.canPlayType('application/vnd.apple.mpegurl')) {
                    video.src = video.getAttribute('data-hls');
                    video.addEventListener('loadedmetadata', function () {
                        video.currentTime = startTime;
                        video.play();
                    });
                    video.addEventListener('timeupdate', function () {
                        if (video.currentTime >= endTime) {
                            video.pause();
                        }
                    });
                }
            });
        });
        """
        soup.body.append(custom_script_tag)

    # Find all quotes in the HTML
    not_found=[]
    for paragraph in soup.find_all('p'):
        paragraph_text = paragraph.get_text()
        # Find all quotes in the paragraph
        quotes = re.findall(r'([\"“”])(.*?)([\"“”])', paragraph_text)
        # Process quotes in reverse order
        for open_quote, quote_text, close_quote in reversed(quotes):
            full_quote = f'{open_quote}{quote_text}{close_quote}'
            clip = find_clip_for_quote(quote_text)
            if clip:
                # Create the appropriate video tag
                if is_m3u8:
                    video_tag = soup.new_tag('video', controls=True, **{'data-hls': video_filepath})
                    video_tag['style'] = "width: 100%; height: auto; margin-bottom: 1em;"
                    video_tag['data-start'] = clip[0]
                    video_tag['data-end'] = clip[1]
                else:
                    video_tag = soup.new_tag('video', controls=True, src=f"{video_filepath}#t={clip[0]},{clip[1]}", crossOrigin="anonymous")
                    video_tag['style'] = "width: 100%; height: auto; margin-bottom: 1em;"

                # Locate the text node containing the full quote
                for content in paragraph.contents:
                    if isinstance(content, str) and full_quote in content:
                        # Split the text node at the quote
                        parts = content.split(full_quote)
                        new_nodes = []
                        for i, part in enumerate(parts):
                            if part:
                                new_nodes.append(part)
                            if i < len(parts) - 1:
                                # Insert video tag and quoted text
                                quote_element = soup.new_tag('span')
                                quote_element.string = full_quote
                                new_nodes.append(video_tag)
                                new_nodes.append(quote_element)
                        # Replace the original text node with new nodes
                        content.replace_with(*new_nodes)
                        break
            else:
                print(f'Quote not found: {quote_text}')
                not_found.append(quote_text)

    # Write the modified HTML content to the output file
    if output_html_filepath:
        with open(output_html_filepath, 'w', encoding='utf-8') as output_html_file:
            output_html_file.write(str(soup))       
            print(f'HTML file generated: {output_html_filepath}')
    else:
        return str(soup)
        
    return not_found
    
def extract_text(content,content_type):
    from bs4 import BeautifulSoup
    from docx import Document
    import fitz
    import io
    
    if 'html' in content_type:
        return BeautifulSoup(content, 'html.parser').get_text()
    elif 'docx' == content_type:
        return "\n".join([paragraph.text for paragraph in Document(io.BytesIO(content)).paragraphs])
    elif 'pdf' == content_type:
        text = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            for page in doc:
                text.append(page.get_text())
        return text
    elif 'txt' == content_type:
        return content.decode('utf-8',errors='replace')
    else:
        raise ValueError("Unsupported file type or content")
        
def load_text_from_path(path):
    with open(path, 'rb') as file:
        content = file.read()
    content_type=path.split(".")[-1]
    return extract_text(content, content_type)

def load_text_from_url(url,timeout=10):
    import requests
    import mimetypes
    from selenium import webdriver
    from selenium.webdriver.firefox.options import Options
    from selenium.webdriver.firefox.service import Service
    from webdriver_manager.firefox import GeckoDriverManager
    
    the_split=url.split(".") #try to split out type
    if the_split[-1] in ['html','docx','text','txt','pdf']: #if likely s static page or file
        headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36'
        }
        response = requests.get(url,headers=headers,timeout=timeout)
        response.raise_for_status() #caller will have to deal with error
        content_type = response.headers['Content-Type'].split(';')[0]
        content_type=(mimetypes.guess_extension(content_type))
        if content_type.startswith('.'):
            content_type=content_type[1:]
        if content_type not in [['html','docx','txt','pdf']]: #if we don't recognize it
            content_type=the_split[-1] #use it from file name
        return extract_text(response.content,content_type)
    else: #if not known to be static page
        firefoxOptions = Options()
        firefoxOptions.add_argument("--headless")
        service = Service(GeckoDriverManager().install())
        driver = webdriver.Firefox(
            options=firefoxOptions,
            service=service,
        )
        driver.set_page_load_timeout(timeout)
        driver.get(url)
        return extract_text(driver.page_source,"html")
            
    

def extract_text_from_path_or_url(path_or_url,content=None,timeout=5):
    import requests
    from bs4 import BeautifulSoup
    from docx import Document
    import fitz
    import io
    import mimetypes
    if not content:
        content=""
    content_type = ""

    if path_or_url.startswith(('http://', 'https://')):
        headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36'
}
        response = requests.get(path_or_url,headers=headers,timeout=timeout)
        response.raise_for_status() #caller will have to deal with error
        content = response.content
        content_type = response.headers['Content-Type'].split(';')[0]
        content_type=(mimetypes.guess_extension(content_type))[1:]
    else:
        if not content: #if not already read
            with open(path_or_url, 'rb') as file:
                content = file.read()
        content_type=path_or_url.split(".")[-1]

    if 'html' in content_type:
        return BeautifulSoup(content, 'html.parser').get_text()
    elif 'docx' == content_type:
        return "\n".join([paragraph.text for paragraph in Document(io.BytesIO(content)).paragraphs])
    elif 'pdf' == content_type:
        text = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            for page in doc:
                text.append(page.get_text())
        return text
    elif 'txt' == content_type:
        return content.decode('utf-8')
    else:
        raise ValueError("Unsupported file type or content")


def extract_text_from_pdf(
        pdf_content: bytes=Field(definition ="binary content of a pdf")):
    #deprecated
    import fitz
    
    with fitz.open(stream=pdf_content, filetype="pdf") as doc:
        text = ""
        for page in doc:
            text += page.get_text()
    return text

    
def extract_text_from_file(file_path):
    #deprecated
    import fitz  # PyMuPDF for PDF files
    from docx import Document  # python-docx for DOCX files
    import html2text  # For HTML files
    
    try:
        if file_path.endswith('.pdf'):
            text = []
            with fitz.open(file_path) as doc:
                for page in doc:
                    text.append(page.get_text())
            return "\n".join(text)
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        elif file_path.endswith('.html'):
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            return html2text.html2text(html_content)
        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
    except Exception as e:
        print(f"Error extracting text from file: {e}")
    return ""


def text_from_web(
        theurl: str=Field(description="url of document to be retrieved"))-> str:
    #deprecated
    import requests
    
    response=requests.get(theurl)
    response.raise_for_status()
    content_type = response.headers['Content-Type']
    if content_type=='application/pdf':
        return extract_text_from_pdf(response.content)
    assert False,f"cannot handle file type {content_type}"
 
class ArxivToolInput(BaseModel):
    query: str = Field(description='A query string.')
    id_list: List[str] =Field (default=[],description='A list of arXiv article IDs to which to limit the search')
    max_results: Optional[int]=Field(default=10,description='The maximum number of results to be returned in an execution of this search')
    sort_by: Optional[Literal["relevance" ,"lastUpdatedDate","submittedDate"]]=Field(default="relevance",description="The sort criterion for results")
    sort_order: Optional[Literal["ascending","descending"]] =Field(default="descending",description="The sort order for results")
 
    
def ArxivTool(input:Annotated[ArxivToolInput,"Input to the search"])->List[Dict[str, Any]]:
#def ArxivTool(input:Annotated[arxiv.Search,"Input to the search"])->List[Dict[str, Any]]:
    from datetime import datetime
    import arxiv    
    fields=['summary','title','published','authors','links']
    print(input)
    input.sort_by=arxiv.SortCriterion(input.sort_by)
    input.sort_order=arxiv.SortOrder(input.sort_order)
    #print(input)
    input_dict=input.model_dump()
    print(input_dict)
    theprompt=arxiv.Search(**input_dict)
    client=arxiv.Client()
    #print(theprompt)
    results=client.results(theprompt)
    all_results=list(results)
    #print(dir(all_results[0]))
    thelist=[]
    for r in all_results: #make a dictionary for each item
        thedict={}
        for field in fields: #looking for each specified field
            try:
                thevalue=getattr(r,field)
                if isinstance(thevalue,datetime):
                    thevalue=thevalue.strftime("%Y-%m-%d")
                elif isinstance(thevalue,list):
                    subfield=""
                    if len(thevalue)>0:
                        if isinstance(thevalue[0],arxiv.Result.Author):
                            subfield='name'
                        elif isinstance(thevalue[0],arxiv.Result.Link):
                            subfield='href'
                        if subfield:
                            thevalue=[getattr(f,subfield) for f in thevalue]
                thedict[field]=thevalue
            except AttributeError:
                pass
        thelist.append(thedict)
    #thelist=[{"summary":r.summary,"title":r.title} for r in all_results]
    #return all_results
    return thelist

class RedditToolInput(BaseModel):
    query: str = Field(description='A query string in lucene format.')
    client_id: Optional[str] = Field(default=None,description="Reddit client id")
    client_secret: Optional[str] = Field(default=None,description="Reddit client secret")
    user_agent: Optional[str] = Field(default="reddit retrieval script",description="user agent for header")
    limit: Optional[int]=Field(default=1,description="maximum number of submissions to return")
    sort: Optional[Literal["relevance","new","hot","comments"]]=Field(default='relevance',description='sort order')
    time_filter: Optional[Literal["all","day","hour","month","week","year"]]=Field(default="year",description="time period for search")
    syntax: Optional[Literal["lucene","cloudsearch","plain"]]=Field(default="lucene",description="type of search")
    

def RedditTool(input:Annotated[RedditToolInput,"Input to the search"])->List[Dict[str, Any]]:
    import os
    from dotenv import load_dotenv
    from datetime import datetime
    import praw
    
    fields=['created_utc','author','title','selftext','link_flare_text','subreddit','url']
    load_dotenv() #just in case
    if (client_id:=input.client_id) is None:
        client_id=os.getenv("REDDIT_CLIENT_ID")
    if (client_secret:=input.client_secret) is None:
        client_secret=os.getenv("REDDIT_CLIENT_SECRET")
    reddit = praw.Reddit(client_id=client_id, 
                         client_secret=client_secret, 
                         user_agent=input.user_agent)
    
    # Perform search and process submissions

    thelist=[]
    for r in reddit.subreddit('all').search(input.query,
        syntax=input.syntax,
        limit=input.limit,
        sort=input.sort,
        time_filter=input.time_filter):
            thedict={}
            for field in fields: #looking for each specified field
                try:
                    thevalue=getattr(r,field)
                    if field=='created_utc':
                        thevalue=datetime.fromtimestamp(thevalue).strftime("%Y-%m-%d")
                        field='created_date'
                    elif field=='author':
                        thevalue=thevalue.name
                    elif field=='subreddit':
                        thevalue=thevalue.display_name
                    thedict[field]=thevalue
                except AttributeError:
                    pass
            thelist.append(thedict)
    return thelist

if __name__ == '__main__': #test code
    """
    import asyncio
    theinput=ArxivToolInput(query="(LLM AND Newsroom) AND submittedDate:[20230101 TO 20240101]")
    answer=ArxivTool(theinput)
    print(len(answer),answer[0])
    loop = asyncio.get_event_loop()
    answer=RedditTool(RedditToolInput(query="LLM in Newsroom",limit=5,time_filter='all'))
    print(len(answer),answer[0])
    """
    while True:
        url=("url: ")
        if not url:
            break
        print(load_text_from_url(url))
